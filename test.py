from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Создаем документ
resume = Document()

# Настройки шрифта по умолчанию
style = resume.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Заголовок с центровкой
header = resume.add_paragraph()
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
header_run = header.add_run('Резюме')
header_run.bold = True
header_run.font.size = Pt(24)

# Имя и контакты
resume.add_paragraph('Имя Фамилия\nPython Backend Developer | Data / Automation\nРоссия (готов к удалённой работе за границей)\nEmail: [твой email] | Телефон: [твой номер] | LinkedIn / GitHub: [ссылки]')

# О себе
resume.add_heading('Обо мне', level=1)
resume.add_paragraph('Python-разработчик с 1,5-летним опытом работы в компании «Транснефть», специализируюсь на обработке данных, автоматизации процессов и разработке бэкенд-сервисов.\n\n- Оптимизировал обработку больших JSON-файлов, превышающих оперативную память.\n- Разрабатывал скрипты для фильтрации и аналитики PostgreSQL.\n- Создавал DAG в Airflow для миграции баз данных MS SQL → PostgreSQL с корректной обработкой форматов и поддержкой обновления больших таблиц (более 1 млн записей).\n- Вне работы занимался созданием Telegram-ботов и изучением бэкенда веб-приложений.\n\nОбладаю аналитическим складом ума, ответственностью, легко нахожу общий язык с коллегами, умею наставлять и объяснять. Владею английским на уровне B2, заинтересован в удалённой работе в международных проектах.')

# Опыт работы
resume.add_heading('Опыт работы', level=1)
resume.add_paragraph('Транснефть — Python Developer\n2024 – 2025\n- Оптимизация обработки больших JSON-файлов для последующей загрузки в базы данных.\n- Создание скриптов для аналитики и фильтрации данных в PostgreSQL.\n- Разработка DAG в Airflow для переноса данных из MS SQL в PostgreSQL, устранение ошибок при миграции, поддержка корректных форматов и постоянного обновления больших таблиц (>1 млн записей).\n- Работа с различными СУБД: PostgreSQL, MS SQL, MySQL, Redis, ClickHouse.\n\nЛичные проекты:\n- Telegram-боты: разработка ботов для автоматизации задач и взаимодействия с пользователями.\n- Изучение бэкенда веб-приложений: проекты с Django и FastAPI, работа с базами данных, обработка данных.')

# Навыки
resume.add_heading('Навыки', level=1)
resume.add_paragraph('Языки программирования: Python (Django, FastAPI, pandas, NumPy)\nБазы данных: PostgreSQL, MS SQL, MySQL, Redis, ClickHouse\nИнструменты: Airflow, Git, Docker (базовые знания), REST API\nАналитика: обработка и фильтрация данных, работа с большими объёмами информации\nСофт-скиллы: аналитическое мышление, работа в команде, наставничество, ответственность')

# Образование
resume.add_heading('Образование', level=1)
resume.add_paragraph('МТУСИ — факультет сетей и систем связи\n3 курс, 2022 – настоящее время\nДополнительно: курсы по Python, Django, FastAPI, анализу данных.')

# Языки
resume.add_heading('Языки', level=1)
resume.add_paragraph('Английский – B2 (разговорный, сертификат)')

# Цель / Позиция
resume.add_heading('Цель / Позиция', level=1)
resume.add_paragraph('Ищу позицию Python Backend Developer, Data Engineer или Automation Developer с возможностью удалённой работы в международной команде для развития навыков в области обработки данных, веб-разработки и автоматизации процессов.')

# Сохраняем документ локально
resume.save('Resume_Python_Backend_Design_Local.docx')

print('Резюме успешно создано: Resume_Python_Backend_Design_Local.docx')
